from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .validate import validate_table
from .util import load_yaml

TABLE_RE = re.compile(r"^\{\{TABLE:\s*(table-\d{2})\s*\}\}$")
FMT_RE = re.compile(r"^<!--\s*fmt:\s*(.*?)\s*-->$")
STYLE_RE = re.compile(r"^<!--\s*style:\s*(.+?)\s*-->$")  # legacy
BLANK_RE = re.compile(r"^\{\{BLANK\}\}$")
PAGE_BREAK_RE = re.compile(r"^\{\{PAGE_BREAK\}\}$")


@dataclass
class Block:
    kind: str  # "heading" | "para" | "table" | "empty" | "page_break"
    text: str = ""
    level: int = 0
    fmt: Dict[str, object] = field(default_factory=dict)
    table_id: Optional[str] = None


def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    # Keep sectPr (section properties) if present, remove everything else
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _set_paragraph_text_with_breaks(paragraph, text: str) -> None:
    paragraph.text = ""  # clear
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            paragraph.add_run(line)
        else:
            r = paragraph.add_run()
            r.add_break()  # line break within paragraph
            r.add_text(line)


def _apply_paragraph_format(paragraph, fmt: Dict[str, object]) -> None:
    # Alignment
    align = fmt.get("align")
    if isinstance(align, str):
        a = align.lower()
        if a == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif a == "right":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif a == "justify":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Run formatting (apply uniformly)
    size = fmt.get("size")
    bold = fmt.get("bold")
    italic = fmt.get("italic")

    for run in paragraph.runs:
        if isinstance(size, (int, float)):
            run.font.size = Pt(float(size))
        if isinstance(bold, bool):
            run.bold = bold
        if isinstance(italic, bool):
            run.italic = italic


def _set_cell_text(cell, text: str) -> None:
    """
    Replace cell contents.

    Newline-separated content is rendered as *separate paragraphs* inside the cell,
    which tends to match Word’s native table-cell structure and rendering more closely.
    """
    cell.text = ""
    parts = str(text).split("\n")
    p0 = cell.paragraphs[0]
    p0.text = parts[0] if parts else ""
    for part in parts[1:]:
        cell.add_paragraph(part)


def _set_repeat_table_header(row) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _parse_fmt_kv(s: str) -> Dict[str, object]:
    """
    Parse key=value pairs where values may be quoted.
    Example: style="List Paragraph" align=center size=24 bold=true
    """
    out: Dict[str, object] = {}

    # Tokenize: key=value where value is "..." or unquoted up to whitespace
    token_re = re.compile(r'(\w+)=(?:"((?:\\.|[^"])*)"|([^\s]+))')
    for m in token_re.finditer(s):
        key = m.group(1)
        if m.group(2) is not None:
            val = m.group(2).replace('\\"', '"')
        else:
            val = m.group(3)

        # Coerce common types
        if isinstance(val, str):
            low = val.lower()
            if low in ("true", "false"):
                out[key] = low == "true"
                continue
            # number?
            try:
                if "." in val:
                    out[key] = float(val)
                else:
                    out[key] = int(val)
                continue
            except Exception:
                pass

        out[key] = val

    return out


def parse_narrative(md_path: Path) -> List[Block]:
    """
    Parse the narrative markdown into a linear sequence of blocks.

    Supported constructs (intentionally small/deterministic):
    - Headings: lines starting with '#', '##', ...
    - Paragraph blocks: one or more lines, separated by blank lines
    - Formatting directive applying to the *next* paragraph: <!-- fmt: ... -->
      (also supports legacy <!-- style: ... -->)
    - Table placeholder: {{TABLE: table-XX}}
    - Explicit empty paragraph: {{BLANK}}
    - Explicit page break: {{PAGE_BREAK}}
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()

    blocks: List[Block] = []
    pending_fmt: Dict[str, object] = {}
    para_lines: List[str] = []
    para_fmt: Optional[Dict[str, object]] = None

    def flush_para() -> None:
        nonlocal para_lines, para_fmt
        if para_lines:
            blocks.append(
                Block(kind="para", text="\n".join(para_lines), fmt=para_fmt or {})
            )
            para_lines = []
            para_fmt = None

    for raw in lines:
        line = raw.rstrip("\n")

        # Blank line = paragraph boundary.
        # Multiple blanks do not create empty paras; use {{BLANK}} for that.
        if line.strip() == "":
            flush_para()
            continue

        # Formatting directive
        m = FMT_RE.match(line.strip())
        if m and not para_lines:
            pending_fmt.update(_parse_fmt_kv(m.group(1)))
            continue

        # Legacy style directive
        m = STYLE_RE.match(line.strip())
        if m and not para_lines:
            pending_fmt["style"] = m.group(1).strip()
            continue

        # Explicit empty paragraph marker
        if BLANK_RE.match(line.strip()):
            flush_para()
            blocks.append(Block(kind="empty"))
            pending_fmt = {}
            continue

        # Explicit page break marker
        if PAGE_BREAK_RE.match(line.strip()):
            flush_para()
            blocks.append(Block(kind="page_break"))
            pending_fmt = {}
            continue

        # Table placeholder
        m = TABLE_RE.match(line.strip())
        if m:
            flush_para()
            blocks.append(Block(kind="table", table_id=m.group(1)))
            pending_fmt = {}
            continue

        # Heading
        if line.startswith("#"):
            flush_para()
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].lstrip()
            blocks.append(Block(kind="heading", level=level, text=text))
            pending_fmt = {}
            continue

        # Regular paragraph line (may be multi-line paragraph)
        if not para_lines:
            para_fmt = dict(pending_fmt)
            pending_fmt = {}
        para_lines.append(line)

    flush_para()
    return blocks


def build_docx(
    template_docx: Path,
    narrative_md: Path,
    tables_dir: Path,
    schemas_dir: Path,
    out_docx: Path,
) -> None:
    common_schema = schemas_dir / "table_common.schema.json"

    doc = Document(str(template_docx))
    _clear_document_body(doc)

    blocks = parse_narrative(narrative_md)

    for b in blocks:
        if b.kind == "heading":
            style = f"Heading {b.level}" if b.level <= 9 else "Heading 9"
            p = doc.add_paragraph("", style=style)
            _set_paragraph_text_with_breaks(p, b.text)
            # allow alignment override on headings if ever needed
            if b.fmt:
                _apply_paragraph_format(p, b.fmt)

        elif b.kind == "para":
            style = str(b.fmt.get("style")) if "style" in b.fmt else "Normal"
            p = doc.add_paragraph("", style=style)
            _set_paragraph_text_with_breaks(p, b.text)
            if b.fmt:
                _apply_paragraph_format(p, b.fmt)

        elif b.kind == "empty":
            doc.add_paragraph("", style="Normal")

        elif b.kind == "page_break":
            # Match Word's common representation: a break inside a paragraph
            p = doc.add_paragraph("", style="Normal")
            r = p.add_run()
            r.add_break(WD_BREAK.PAGE)

        elif b.kind == "table":
            assert b.table_id
            ypath = tables_dir / f"{b.table_id}.yaml"
            spath = schemas_dir / f"{b.table_id}.schema.json"
            if not ypath.exists():
                raise FileNotFoundError(f"Missing table YAML: {ypath}")
            if not spath.exists():
                raise FileNotFoundError(f"Missing table schema: {spath}")

            validate_table(ypath, spath, common_schema)
            table_data = load_yaml(ypath)

            columns = table_data["columns"]
            rows = table_data["rows"]
            style_name = table_data.get("style") or "Table Grid"

            tbl = doc.add_table(rows=len(rows) + 1, cols=len(columns))
            tbl.style = style_name

            # Header row
            header_row = tbl.rows[0]
            _set_repeat_table_header(header_row)
            for c_idx, col in enumerate(columns):
                cell = header_row.cells[c_idx]
                _set_cell_text(cell, str(col["title"]))
                # Bold header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Data rows
            for r_idx, row_obj in enumerate(rows, start=1):
                tr = tbl.rows[r_idx]
                for c_idx, col in enumerate(columns):
                    key = col["key"]
                    txt = row_obj.get(key, "")
                    _set_cell_text(tr.cells[c_idx], str(txt))

        else:
            raise ValueError(f"Unknown block kind: {b.kind}")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))

from __future__ import annotations

import difflib
from dataclasses import dataclass
from pathlib import Path
from typing import List, Union

from docx import Document


@dataclass
class ParaRep:
    style: str
    text: str


@dataclass
class TableRep:
    style: str
    rows: List[List[str]]  # includes header row


ElementRep = Union[ParaRep, TableRep]


def _flatten_docx(docx_path: Path) -> List[ElementRep]:
    doc = Document(str(docx_path))
    p_map = {p._p: p for p in doc.paragraphs}
    t_map = {t._tbl: t for t in doc.tables}

    elems: List[ElementRep] = []
    for child in doc._element.body.iterchildren():
        tag = child.tag
        if tag.endswith("}sectPr"):
            continue
        if tag.endswith("}p"):
            p = p_map.get(child)
            if p is None:
                continue
            elems.append(ParaRep(style=p.style.name, text=p.text.rstrip("\n")))
        elif tag.endswith("}tbl"):
            t = t_map.get(child)
            if t is None:
                continue
            matrix = []
            for row in t.rows:
                matrix.append([cell.text.strip() for cell in row.cells])
            elems.append(TableRep(style=t.style.name if t.style else "", rows=matrix))
    return elems


def compare_docx_text(
    baseline_docx: Path, generated_docx: Path, max_mismatches: int = 80
) -> str:
    base = _flatten_docx(baseline_docx)
    gen = _flatten_docx(generated_docx)

    lines: List[str] = []
    lines.append("# DOCX similarity report")
    lines.append("")
    lines.append(f"- Baseline: `{baseline_docx.name}`")
    lines.append(f"- Generated: `{generated_docx.name}`")
    lines.append("")
    lines.append("## Structural counts")
    lines.append("")

    base_para_count = sum(isinstance(x, ParaRep) for x in base)
    gen_para_count = sum(isinstance(x, ParaRep) for x in gen)
    base_table_count = sum(isinstance(x, TableRep) for x in base)
    gen_table_count = sum(isinstance(x, TableRep) for x in gen)

    lines.append(
        f"- Elements: baseline **{len(base)}**, generated **{len(gen)}** "
        "(paragraphs+tables)"
    )
    lines.append(
        f"- Paragraphs: baseline **{base_para_count}**, generated **{gen_para_count}**"
    )
    lines.append(
        f"- Tables: baseline **{base_table_count}**, generated **{gen_table_count}**"
    )
    lines.append("")

    # Compare sequence element-wise up to min length
    mismatches: List[str] = []
    min_len = min(len(base), len(gen))
    for i in range(min_len):
        b = base[i]
        g = gen[i]
        if b.__class__ is not g.__class__:
            mismatches.append(
                f"- Element {i}: type differs "
                f"(baseline {type(b).__name__}, generated {type(g).__name__})"
            )
            continue
        if isinstance(b, ParaRep) and isinstance(g, ParaRep):
            if b.text != g.text or b.style != g.style:
                mismatches.append(
                    f"- Paragraph {i}: style/text mismatch\n"
                    f"  - baseline style=`{b.style}` text=`{b.text[:120]}`\n"
                    f"  - generated style=`{g.style}` text=`{g.text[:120]}`"
                )
        else:
            assert isinstance(b, TableRep) and isinstance(g, TableRep)
            if b.rows != g.rows or b.style != g.style:
                baseline_shape = f"{len(b.rows)}x{len(b.rows[0]) if b.rows else 0}"
                generated_shape = f"{len(g.rows)}x{len(g.rows[0]) if g.rows else 0}"
                mismatches.append(
                    f"- Table {i}: style/content mismatch\n"
                    f"  - baseline style=`{b.style}` rows={baseline_shape}\n"
                    f"  - generated style=`{g.style}` rows={generated_shape}"
                )

    if len(base) != len(gen):
        mismatches.append(
            f"- Document length differs: baseline {len(base)} vs generated {len(gen)}"
        )

    lines.append("## Element-by-element mismatches")
    lines.append("")
    if not mismatches:
        lines.append(
            "✅ No mismatches detected in paragraph text / table cell text, "
            "or element ordering."
        )
    else:
        lines.append(
            f"Found **{len(mismatches)}** mismatches (showing up to {max_mismatches})."
        )
        lines.append("")
        lines.extend(mismatches[:max_mismatches])
        if len(mismatches) > max_mismatches:
            lines.append(f"\n… plus {len(mismatches) - max_mismatches} more.")

    # Overall text similarity (coarse)
    base_text = "\n".join(x.text for x in base if isinstance(x, ParaRep))
    gen_text = "\n".join(x.text for x in gen if isinstance(x, ParaRep))
    ratio = difflib.SequenceMatcher(a=base_text, b=gen_text).ratio()

    lines.append("")
    lines.append("## Coarse narrative-text similarity")
    lines.append("")
    lines.append(f"- SequenceMatcher ratio (paragraph text only): **{ratio:.4f}**")
    lines.append("")
    if ratio < 0.98:
        # include a small diff snippet
        diff = difflib.unified_diff(
            base_text.splitlines(),
            gen_text.splitlines(),
            fromfile="baseline",
            tofile="generated",
            lineterm="",
        )
        snippet = []
        for j, dline in enumerate(diff):
            if j > 200:
                snippet.append("… (diff truncated)")
                break
            snippet.append(dline)
        lines.append("### Diff snippet (first ~200 lines)")
        lines.append("")
        lines.append("```diff")
        lines.extend(snippet)
        lines.append("```")

    return "\n".join(lines) + "\n"
